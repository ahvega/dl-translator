from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def _set_run_font_arial(run) -> None:
    run.font.name = "Arial"
    rfonts = run._element.rPr.rFonts if run._element.rPr is not None else None
    if rfonts is None:
        run._element.get_or_add_rPr().get_or_add_rFonts()
        rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:eastAsia"), "Arial")
    rfonts.set(qn("w:cs"), "Arial")


def _set_paragraphs_font_arial(paragraphs) -> None:
    for paragraph in paragraphs:
        for run in paragraph.runs:
            _set_run_font_arial(run)


def _enforce_arial_font(output_path: Path) -> None:
    doc = Document(output_path)

    if "Normal" in doc.styles:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        rpr = normal._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Arial")
        rfonts.set(qn("w:hAnsi"), "Arial")
        rfonts.set(qn("w:eastAsia"), "Arial")
        rfonts.set(qn("w:cs"), "Arial")

    _set_paragraphs_font_arial(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _set_paragraphs_font_arial(cell.paragraphs)

    for section in doc.sections:
        _set_paragraphs_font_arial(section.header.paragraphs)
        _set_paragraphs_font_arial(section.footer.paragraphs)

    doc.save(output_path)


def markdown_to_docx(
    md_text: str,
    output_path: Path,
    resource_parent: Path | None = None,
    author: str | None = None,
) -> None:
    """Write DOCX from Markdown using Pandoc (pypandoc)."""
    import pypandoc

    extra_args: list[str] = []
    if resource_parent is not None and resource_parent.is_dir():
        extra_args.append(f"--resource-path={resource_parent}")
    if author:
        extra_args.extend(["--metadata", f"author={author}"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        pypandoc.convert_text(
            md_text,
            "docx",
            format="md",
            outputfile=str(output_path),
            extra_args=extra_args,
        )
    except (OSError, RuntimeError) as e:
        raise RuntimeError(
            "DOCX export requires Pandoc installed and on PATH. "
            "See https://pandoc.org/installing.html (Windows: winget install pandoc)."
        ) from e

    _enforce_arial_font(output_path)
