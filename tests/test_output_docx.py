from pathlib import Path

from docx import Document

from dl_translator import cli
from dl_translator.output_docx import _enforce_arial_font


def test_convert_markdown_outputs_to_docx_continues_after_failure(monkeypatch, tmp_path):
    first = tmp_path / "one.md"
    second = tmp_path / "two.md"
    first.write_text("one", encoding="utf-8")
    second.write_text("two", encoding="utf-8")

    calls: list[Path] = []

    def fake_markdown_to_docx(md_text, output_path, resource_parent=None):
        calls.append(output_path)
        if output_path.name == "one.docx":
            raise RuntimeError("pandoc missing")

    monkeypatch.setattr(cli, "markdown_to_docx", fake_markdown_to_docx)

    failures = cli._convert_markdown_outputs_to_docx(
        [(first, tmp_path), (second, tmp_path)]
    )

    assert failures == 1
    assert calls == [tmp_path / "one.docx", tmp_path / "two.docx"]


def test_enforce_arial_font_updates_existing_runs(tmp_path):
    path = tmp_path / "sample.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Hello")
    run.font.name = "Times New Roman"
    doc.save(path)

    _enforce_arial_font(path)

    updated = Document(path)
    assert updated.paragraphs[0].runs[0].font.name == "Arial"
